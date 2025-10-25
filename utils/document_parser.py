"""
Module d'extraction de contenu depuis fichiers Word et PDF
Gère le texte et les images pour alimenter Claude
VERSION OPTIMISÉE pour extraction rapide
"""

import io
import base64
from typing import List, Dict, Tuple
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import fitz  # PyMuPDF
from PIL import Image


class DocumentParser:
    """Classe pour extraire texte et images de documents Word/PDF"""
    
    # Configuration optimisée
    MAX_IMAGES = 10  # Limite d'images à extraire
    MAX_IMAGE_SIZE = (1024, 1024)  # Résolution max (largeur, hauteur)
    IMAGE_QUALITY = 85  # Qualité JPEG (85 = bon compromis qualité/taille)
    
    @staticmethod
    def extract_from_word(file_bytes: bytes) -> Tuple[str, List[Dict]]:
        """
        Extrait le texte et les images d'un fichier Word
        
        Args:
            file_bytes: Contenu du fichier Word en bytes
            
        Returns:
            Tuple (texte_complet, liste_images)
            - texte_complet: String contenant tout le texte
            - liste_images: Liste de dicts avec {data: base64, format: str}
        """
        doc = Document(io.BytesIO(file_bytes))
        
        # Extraction du texte
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extraction des tableaux (souvent utilisés en médecine)
        for table in doc.tables:
            table_text = DocumentParser._extract_table_text(table)
            if table_text:
                text_parts.append(table_text)
        
        full_text = "\n\n".join(text_parts)
        
        # Extraction des images (OPTIMISÉ - limite à MAX_IMAGES)
        images = []
        image_count = 0
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref and image_count < DocumentParser.MAX_IMAGES:
                try:
                    image_data = rel.target_part.blob
                    
                    # Optimiser l'image avant conversion base64
                    optimized_data = DocumentParser._optimize_image(image_data)
                    
                    # Convertir en base64
                    img_base64 = base64.b64encode(optimized_data).decode('utf-8')
                    
                    # Déterminer le format
                    img_format = DocumentParser._get_image_format(optimized_data)
                    
                    images.append({
                        'data': img_base64,
                        'format': img_format
                    })
                    
                    image_count += 1
                    
                except Exception as e:
                    print(f"Erreur extraction image: {e}")
                    continue
        
        return full_text, images
    
    @staticmethod
    def _extract_table_text(table: Table) -> str:
        """Extrait le texte d'un tableau Word"""
        table_lines = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_lines.append(row_text)
        return "\n".join(table_lines)
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> Tuple[str, List[Dict]]:
        """
        Extrait le texte et les images d'un fichier PDF
        VERSION OPTIMISÉE - Limite extraction et compresse images
        
        Args:
            file_bytes: Contenu du fichier PDF en bytes
            
        Returns:
            Tuple (texte_complet, liste_images)
        """
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        text_parts = []
        images = []
        image_count = 0
        
        for page_num in range(len(doc)):
            # Arrêter si on a déjà assez d'images
            if image_count >= DocumentParser.MAX_IMAGES:
                # Continuer à extraire le texte mais pas les images
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                continue
            
            page = doc[page_num]
            
            # Extraction du texte
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            # Extraction des images (LIMITÉE)
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                if image_count >= DocumentParser.MAX_IMAGES:
                    break
                
                try:
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # OPTIMISATION : Compresser l'image
                    optimized_bytes = DocumentParser._optimize_image(image_bytes)
                    
                    # Convertir en base64
                    img_base64 = base64.b64encode(optimized_bytes).decode('utf-8')
                    img_format = base_image["ext"]
                    
                    images.append({
                        'data': img_base64,
                        'format': img_format
                    })
                    
                    image_count += 1
                    
                except Exception as e:
                    print(f"Erreur extraction image PDF page {page_num + 1}: {e}")
                    continue
        
        full_text = "\n\n".join(text_parts)
        doc.close()
        
        return full_text, images
    
    @staticmethod
    def _optimize_image(image_bytes: bytes) -> bytes:
        """
        Optimise une image (redimensionne et compresse)
        
        Args:
            image_bytes: Image originale en bytes
            
        Returns:
            Image optimisée en bytes
        """
        try:
            img = Image.open(io.BytesIO(image_bytes))
            
            # Convertir en RGB si nécessaire (pour JPEG)
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Redimensionner si trop grande
            if img.size[0] > DocumentParser.MAX_IMAGE_SIZE[0] or img.size[1] > DocumentParser.MAX_IMAGE_SIZE[1]:
                img.thumbnail(DocumentParser.MAX_IMAGE_SIZE, Image.Resampling.LANCZOS)
            
            # Compresser en JPEG
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=DocumentParser.IMAGE_QUALITY, optimize=True)
            
            return output.getvalue()
            
        except Exception as e:
            print(f"Erreur optimisation image: {e}")
            # Si erreur, retourner l'original
            return image_bytes
    
    @staticmethod
    def _get_image_format(image_bytes: bytes) -> str:
        """Détermine le format d'une image depuis ses bytes"""
        try:
            img = Image.open(io.BytesIO(image_bytes))
            format_map = {
                'JPEG': 'jpeg',
                'PNG': 'png',
                'GIF': 'gif',
                'WEBP': 'webp'
            }
            return format_map.get(img.format, 'jpeg')
        except:
            return 'jpeg'  # Format par défaut
    
    @staticmethod
    def parse_document(file_bytes: bytes, file_type: str) -> Tuple[str, List[Dict]]:
        """
        Point d'entrée principal pour parser un document
        
        Args:
            file_bytes: Contenu du fichier
            file_type: 'docx' ou 'pdf'
            
        Returns:
            Tuple (texte, images) avec images optimisées et limitées
        """
        if file_type == 'docx':
            return DocumentParser.extract_from_word(file_bytes)
        elif file_type == 'pdf':
            return DocumentParser.extract_from_pdf(file_bytes)
        else:
            raise ValueError(f"Type de fichier non supporté: {file_type}")